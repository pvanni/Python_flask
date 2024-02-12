import subprocess
from docx.shared import Inches
from docx import Document
import pandas as pd
import numpy as np


class ShellCommand:
    """Takes shell command and run it"""

    def __init__(self, comstring):
        self.runcommand(comstring)

    def runcommand(self, command):
        process = subprocess.Popen(command.split('\n'), stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=True)

        process.wait()

        stdout, stderr = process.communicate()

        if process.returncode > 0:
            print(f"{stdout.decode('utf-8')}")
            print(f"{stderr.decode('utf-8')}")
            raise Exception('Error')

        return stdout


def results_manager(target_variables=False):
    """Decorator factory to write result txt files and iterate through all target variables"""
    def results_decorator(func):
        def results_wrapper(*args, **kwargs):

            if target_variables:
                para = func(*args, **kwargs)
                if para:
                    with open(f"{kwargs['project_path']}/methods/{func.__name__}.txt", 'w+') as f:
                        for line in para:
                            f.write(f'{line}\n')

            else:
                # This is a method function, so write the paragraphs into a methods folder
                para = func(*args, **kwargs)
                with open(f"{kwargs['project_path']}/methods/{func.__name__}.txt", 'w+') as f:
                    for line in para:
                        f.write(f'{line}\n')

        return results_wrapper

    return results_decorator


def report_writer(document, heading, paragraphs, pictures=None):
    """Take the document and expand it with the paragraphs in the list"""

    document.add_heading(heading, level=1)
    for para in paragraphs:
        document.add_paragraph(para)

    if pictures:
        for pic in pictures:
            document.add_picture(pic, width=Inches(6.0))
